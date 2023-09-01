import re
import sys

import docx
from flask import Flask, render_template, request

app = Flask(__name__)

# Đường dẫn đến tệp .docx
docx_file_path = "bieumau/hồ sơ mời sơ tuyển xây lắp qua mạng.docx"


def find_matches_in_text(text, pattern):
    matches = re.findall(pattern, text)
    return [match[1] if match[1] != "" else match[0].strip() for match in matches]

def replace_multiple_underscores(match):
    return re.sub(r'_+', '___', match)


@app.route("/", methods=["GET", "POST"])
def index():
    # Đọc tệp .docx
    doc = docx.Document(docx_file_path)

    # Định nghĩa pattern
    pattern = r"([^\n]+_{3,}[^\n]+|([^\n]+)\n\s+_{3,})"
    document_text = "\n".join([para.text for para in doc.paragraphs])
    matches = find_matches_in_text(document_text, pattern)

    replaced_matches = [replace_multiple_underscores(match) for match in matches]

    # Trích xuất thông tin từ các bảng trong tệp .docx
    tables_data = []
    table_names = []  # Danh sách tên bảng
    current_table_name = None
    for table in doc.tables:
        table_data = []
        for i, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if not cell_text:
                    cell_text = '<input type="text">'
                else:
                    # Kiểm tra và thay thế dấu _ liền kề bằng input text nếu cần
                    cell_text = re.sub(r'_+', '<input type="text">', cell_text)
                row_data.append(cell_text)
            table_data.append(row_data)
        tables_data.append(table_data)
    if current_table_name:
        table_names.append(current_table_name)  # Lưu tên bảng cuối cùng vào danh sách

    if request.method == "POST":
        edited_text = request.form.get("edited_text", "")
        updated_document_text = re.sub(pattern, edited_text, document_text)
        return "Changes saved successfully."

    return render_template("result.html", matches=replaced_matches, tables_data=tables_data, table_names=table_names)


if __name__ == "__main__":
    app.run(debug=True)
